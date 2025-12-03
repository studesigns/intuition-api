#!/usr/bin/env python3
"""
Create Word documents for hallucination prevention testing
Generates 3 comprehensive policy documents with strategic testing features
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def add_heading_formatted(doc, text, level, bold=True):
    """Add a formatted heading"""
    heading = doc.add_heading(text, level=level)
    if bold:
        for run in heading.runs:
            run.font.bold = True
            run.font.size = Pt(14 if level == 1 else 12)

def add_section_heading(doc, number, title):
    """Add a section heading with number"""
    p = doc.add_paragraph()
    p.style = 'Heading 2'
    run = p.add_run(f"{number}. {title}")
    run.font.bold = True
    run.font.size = Pt(12)

def create_document_1():
    """Create Global Entertainment & Client Relations Policy"""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("ZENITH CORPORATION")
    title_run.font.bold = True
    title_run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("Global Entertainment & Client Relations Policy")
    subtitle_run.font.bold = True
    subtitle_run.font.size = Pt(14)

    # Header info
    doc.add_paragraph("Effective Date: January 1, 2025")
    doc.add_paragraph("Policy Number: ZEN-CLP-2025-01")
    doc.add_paragraph("Classification: GLOBAL - APPLIES TO ALL REGIONS")
    doc.add_paragraph()

    # Executive Summary
    add_section_heading(doc, "EXECUTIVE SUMMARY", "")
    doc.add_paragraph(
        "This policy establishes standards for employee engagement with clients in social and entertainment settings. "
        "All Zenith employees are authorized to conduct client entertainment within the parameters established by this policy. "
        "This policy applies to all geographic regions worldwide unless superseded by regional addendums."
    )

    # Section 1
    add_section_heading(doc, "1", "SCOPE AND APPLICABILITY")
    doc.add_paragraph("Section 1.1 Geographic Scope: GLOBAL")
    doc.add_paragraph(
        "This policy applies to all employees, contractors, and representatives of Zenith Corporation worldwide, including:\n"
        "- North America (USA, Canada, Mexico)\n"
        "- Europe (EU and non-EU countries)\n"
        "- Asia-Pacific Region\n"
        "- Middle East\n"
        "- Africa\n"
        "- All other geographies",
        style='List Bullet'
    )

    doc.add_paragraph("Section 1.2 Authorization Level: Standard employee approval authority")
    doc.add_paragraph(
        "- Employees may approve entertainment activities up to USD 500 per event\n"
        "- Managers may approve up to USD 1,500 per event\n"
        "- Director approval required for USD 1,500-5,000\n"
        "- VP approval required for USD 5,000+",
        style='List Bullet'
    )

    # Section 2
    add_section_heading(doc, "2", "PERMITTED ENTERTAINMENT ACTIVITIES")
    doc.add_paragraph(
        "The following entertainment activities are explicitly PERMITTED under this global policy:"
    )

    doc.add_paragraph("Section 2.1 Dining & Beverages")
    doc.add_paragraph(
        "- Restaurant meals (any cuisine type)\n"
        "- Cocktail bars and lounges\n"
        "- Wine and spirits tastings\n"
        "- Coffee and tea venues\n"
        "- Hotel dining experiences",
        style='List Bullet'
    )

    doc.add_paragraph("Section 2.2 Sports & Recreation")
    doc.add_paragraph(
        "- Golf outings\n"
        "- Tennis facilities\n"
        "- Basketball courts\n"
        "- Swimming facilities\n"
        "- Ski resorts\n"
        "- Water sports (boating, sailing)",
        style='List Bullet'
    )

    doc.add_paragraph("Section 2.3 Cultural Activities")
    doc.add_paragraph(
        "- Theater and performing arts\n"
        "- Museums and galleries\n"
        "- Movie theaters\n"
        "- Concert venues\n"
        "- Historical site tours\n"
        "- Art festivals",
        style='List Bullet'
    )

    doc.add_paragraph("Section 2.4 Business Meals")
    doc.add_paragraph(
        "- Breakfast meetings\n"
        "- Lunch meetings\n"
        "- Dinner meetings\n"
        "- Team lunches\n"
        "- Client appreciation events",
        style='List Bullet'
    )

    # Section 3
    add_section_heading(doc, "3", "RESTRICTED ACTIVITIES (REQUIRES SPECIAL APPROVAL)")
    doc.add_paragraph(
        "The following activities are NOT automatically approved and require explicit manager/director approval:"
    )

    doc.add_paragraph("Section 3.1 High-Cost Events")
    doc.add_paragraph(
        "- Events exceeding USD 500 per person require manager pre-approval\n"
        "- Events exceeding USD 1,500 per person require VP approval\n"
        "- Luxury resort entertainment requires budget approval",
        style='List Bullet'
    )

    doc.add_paragraph("Section 3.2 Overnight Activities")
    doc.add_paragraph(
        "- Multi-day entertainment events\n"
        "- Resort stays beyond standard business travel\n"
        "- Adventure tourism activities",
        style='List Bullet'
    )

    # Section 4 - Expenses
    add_section_heading(doc, "4", "EXPENSES AND REIMBURSEMENT")
    doc.add_paragraph("Section 4.1 Reimbursable Expenses")
    doc.add_paragraph(
        "- Meal costs: 100% reimbursable per policy limits\n"
        "- Tickets and entrance fees: 100% reimbursable\n"
        "- Transportation to entertainment: 100% reimbursable\n"
        "- Gratuities: 20% of pre-tax total (standard rates)",
        style='List Bullet'
    )

    # Documentation
    add_section_heading(doc, "5", "DOCUMENTATION REQUIREMENTS")
    doc.add_paragraph("All entertainment expenses must include:")
    doc.add_paragraph(
        "- Date of activity\n"
        "- Client name and organization\n"
        "- Business purpose\n"
        "- Amount spent\n"
        "- Attendees (names and roles)\n"
        "- Receipt or invoice",
        style='List Bullet'
    )

    # Signature block
    doc.add_paragraph()
    doc.add_paragraph("Policy Approval:")
    doc.add_paragraph("Global Chief Compliance Officer: Sarah Mitchell")
    doc.add_paragraph("Global CFO: Robert Chen")
    doc.add_paragraph("CEO: Margaret Williams")

    doc.add_paragraph()
    doc.add_paragraph("Last Revised: January 1, 2025")

    return doc

def create_document_2():
    """Create Asia-Pacific Regional Addendum"""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("ZENITH CORPORATION")
    title_run.font.bold = True
    title_run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("Asia-Pacific Region: Prohibited High-Risk Entertainment Activities")
    subtitle_run.font.bold = True
    subtitle_run.font.size = Pt(12)

    subsubtitle = doc.add_paragraph()
    subsubtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subsubtitle_run = subsubtitle.add_run("Regional Addendum to Global Entertainment Policy")
    subsubtitle_run.font.italic = True
    subsubtitle_run.font.size = Pt(11)

    # Header info
    doc.add_paragraph("Effective Date: January 1, 2025")
    doc.add_paragraph("Policy Number: ZEN-CLP-APAC-2025-01")
    doc.add_paragraph("Geographic Scope: ASIA-PACIFIC REGION ONLY")
    doc.add_paragraph()

    # Executive Summary
    add_section_heading(doc, "EXECUTIVE SUMMARY", "")
    p = doc.add_paragraph(
        "This addendum modifies the Global Entertainment & Client Relations Policy specifically for the Asia-Pacific region. "
        "Due to regulatory, cultural, and business-specific considerations in APAC, certain entertainment activities are "
        "PROHIBITED in this region, even though they may be permitted globally."
    )

    # CRITICAL - SCOPE STATEMENT
    critical = doc.add_paragraph()
    critical_run = critical.add_run(
        "CRITICAL: This addendum applies ONLY to the following APAC countries:\n"
        "- China (PRC)\n"
        "- Japan\n"
        "- South Korea\n"
        "- Taiwan\n"
        "- Vietnam\n"
        "- Indonesia\n"
        "- Thailand\n"
        "- Malaysia\n"
        "- Philippines\n"
        "- Singapore"
    )
    critical_run.font.bold = True
    critical_run.font.color.rgb = RGBColor(255, 0, 0)

    doc.add_paragraph()
    not_applies = doc.add_paragraph()
    not_applies_run = not_applies.add_run(
        "This addendum does NOT apply to: Europe, North America, Middle East, Africa, or any other non-APAC regions."
    )
    not_applies_run.font.bold = True

    doc.add_paragraph()

    # Section 1 - Prohibited Activities
    add_section_heading(doc, "1", "PROHIBITED ENTERTAINMENT ACTIVITIES IN APAC REGION")
    doc.add_paragraph("Section 1.1 Explicitly Prohibited Activities")
    doc.add_paragraph(
        "The following entertainment activities are STRICTLY PROHIBITED in the Asia-Pacific region:"
    )

    # Karaoke
    doc.add_paragraph("Section 1.1.1 Karaoke Venues")
    karaoke_p = doc.add_paragraph(
        "- Private karaoke bars\n"
        "- Karaoke lounges\n"
        "- Entertainment complexes featuring karaoke\n"
        "- Establishment of any type where karaoke is primary activity",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Reason for Prohibition: APAC regulatory environment. Certain karaoke establishments in the region operate under "
        "questionable legal structures and may involve improper inducements. Zenith employees must avoid any appearance of impropriety."
    )
    doc.add_paragraph(
        "Prohibited in: China, Japan, South Korea, Vietnam, Indonesia, Thailand, Malaysia"
    )
    doc.add_paragraph(
        "NOT Prohibited in: Europe, Americas, Middle East, Africa"
    )

    # Nightclub
    doc.add_paragraph("Section 1.1.2 Nightclub Entertainment (Late-Night)")
    doc.add_paragraph(
        "- Nightclubs and dance clubs\n"
        "- Discotheques with adult entertainment\n"
        "- Late-night entertainment venues (operating after 11 PM as primary business)",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Reason for Prohibition: Regulatory risk and business conduct standards in APAC."
    )
    doc.add_paragraph(
        "Prohibited in: China, Japan, South Korea, Vietnam, Indonesia, Thailand"
    )

    # Hostess Bars
    doc.add_paragraph("Section 1.1.3 Hostess Bars and Adult Entertainment Establishments")
    doc.add_paragraph(
        "- Any venue with hostess services\n"
        "- Adult entertainment venues\n"
        "- Escort service coordination\n"
        "- Any activity involving hired companions for entertainment",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Reason for Prohibition: Anti-bribery and improper inducement policies. APAC regulatory environment has "
        "stricter requirements around this."
    )

    # Gambling
    doc.add_paragraph("Section 1.1.4 Gambling Facilities")
    doc.add_paragraph(
        "- Casinos\n"
        "- Gambling establishments\n"
        "- Gaming venues\n"
        "- Betting parlors",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Reason for Prohibition: Compliance with FCPA and regional anti-corruption standards."
    )
    doc.add_paragraph(
        "Prohibited in: China, Vietnam, Indonesia, Malaysia"
    )

    # Permitted Activities
    doc.add_paragraph()
    add_section_heading(doc, "2", "PERMITTED ACTIVITIES IN APAC REGION")
    doc.add_paragraph(
        "The following activities ARE permitted in APAC (per Global Policy):"
    )
    doc.add_paragraph(
        "- Fine dining restaurants\n"
        "- Golf outings (courses approved by Finance)\n"
        "- Museum visits and cultural events\n"
        "- Theater and performing arts\n"
        "- Hotel business dining\n"
        "- Team activities (sports, recreational facilities)\n"
        "- Business conferences and seminars",
        style='List Bullet'
    )

    # Violations
    doc.add_paragraph()
    add_section_heading(doc, "3", "VIOLATIONS IN APAC REGION")
    doc.add_paragraph("Section 3.1 APAC-Specific Enforcement")
    doc.add_paragraph(
        "Violations of this APAC addendum carry enhanced penalties:\n"
        "- First violation: Mandatory retraining + written warning + loss of entertainment authority\n"
        "- Second violation: Suspension of business entertainment privileges + mandatory investigation\n"
        "- Third violation: Disciplinary action up to immediate termination",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Reason: Heightened compliance risk in region requires strict enforcement."
    )

    # Scope Reminder
    doc.add_paragraph()
    scope_reminder = doc.add_paragraph()
    scope_reminder_run = scope_reminder.add_run(
        "SCOPE REMINDER: This addendum applies ONLY to Asia-Pacific countries: "
        "China, Japan, South Korea, Taiwan, Vietnam, Indonesia, Thailand, Malaysia, Philippines, Singapore"
    )
    scope_reminder_run.font.bold = True
    scope_reminder_run.font.color.rgb = RGBColor(255, 0, 0)

    # Approval
    doc.add_paragraph()
    doc.add_paragraph("Approval Authority:")
    doc.add_paragraph("APAC Regional President: David Liu")
    doc.add_paragraph("APAC Regional Compliance Officer: Jennifer Wong")
    doc.add_paragraph("Global Chief Compliance Officer: Sarah Mitchell")

    return doc

def create_document_3():
    """Create Global Business Travel & Entertainment Expenses Policy"""
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("ZENITH CORPORATION")
    title_run.font.bold = True
    title_run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.add_run("Global Business Travel & Entertainment Expenses Policy")
    subtitle_run.font.bold = True
    subtitle_run.font.size = Pt(14)

    # Header info
    doc.add_paragraph("Policy Number: ZEN-TRAVEL-2025-01")
    doc.add_paragraph("Classification: GLOBAL - APPLIES TO ALL REGIONS")
    doc.add_paragraph("Effective Date: January 1, 2025")
    doc.add_paragraph()

    # Executive Summary
    add_section_heading(doc, "EXECUTIVE SUMMARY", "")
    doc.add_paragraph(
        "This policy governs business travel and entertainment expenses for all Zenith Corporation employees worldwide. "
        "This policy works in conjunction with the Global Entertainment & Client Relations Policy and regional addendums."
    )

    # Section 1
    add_section_heading(doc, "1", "ENTERTAINMENT EXPENSE CATEGORIES")
    doc.add_paragraph("Section 1.1 Categories")
    doc.add_paragraph(
        "Entertainment expenses are categorized as:\n"
        "- Client Entertainment: Meals and activities with external clients\n"
        "- Team Building: Internal employee activities\n"
        "- Business Development: Prospecting and networking\n"
        "- Relationship Maintenance: Ongoing client engagement",
        style='List Bullet'
    )

    # Section 2 - IMPORTANT
    add_section_heading(doc, "2", "APPROVAL MATRIX")
    doc.add_paragraph("Section 2.1 Global Approval Requirements")
    doc.add_paragraph(
        "Under USD 100: No approval required\n"
        "USD 100-500: Manager approval required\n"
        "USD 500-1,500: Director approval required\n"
        "USD 1,500-5,000: VP approval required\n"
        "Over USD 5,000: C-Suite approval (CFO or CEO)",
        style='List Bullet'
    )

    doc.add_paragraph("Section 2.2 APAC Region Exception")
    critical_note = doc.add_paragraph(
        "In Asia-Pacific region, the following modifications apply:\n"
        "- Under USD 300 requires Manager approval (vs. no approval globally)\n"
        "- USD 300-1,000 requires VP approval (vs. Manager globally)\n"
        "- ALL APAC expenses require Finance pre-approval regardless of amount"
    )
    critical_note_run = critical_note.runs[0]
    critical_note_run.font.bold = True
    critical_note_run.font.color.rgb = RGBColor(255, 0, 0)

    # Section 3
    add_section_heading(doc, "3", "NON-REIMBURSABLE ITEMS")
    doc.add_paragraph("Section 3.1 Prohibited Expenses")
    doc.add_paragraph(
        "The following are never reimbursable:\n"
        "- Personal entertainment (movies, concerts for self only)\n"
        "- Alcohol for personal consumption (only in business meal context)\n"
        "- Activities that violate regional addendums\n"
        "- Expenses at prohibited venues\n"
        "- Spousal or family member entertainment\n"
        "- Gambling losses",
        style='List Bullet'
    )

    # Section 4
    add_section_heading(doc, "4", "COMPLIANCE WITH ANTI-CORRUPTION LAWS")
    doc.add_paragraph("Section 4.1 FCPA Compliance")
    doc.add_paragraph(
        "All entertainment must comply with:\n"
        "- Foreign Corrupt Practices Act (FCPA)\n"
        "- UK Bribery Act\n"
        "- Local anti-corruption laws\n"
        "- Zenith Code of Conduct",
        style='List Bullet'
    )

    # Section 5 - Integration
    add_section_heading(doc, "5", "REGIONAL POLICY INTEGRATION")
    doc.add_paragraph("Section 5.1 Regional Addendums")
    doc.add_paragraph(
        "Where regional addendums exist, they take precedence:\n"
        "- Asia-Pacific addendum supersedes this policy on prohibited activities\n"
        "- Regional restrictions are MORE restrictive than global policy\n"
        "- Employees must follow most restrictive applicable policy",
        style='List Bullet'
    )

    # CRITICAL EXAMPLE
    example = doc.add_paragraph()
    example_run = example.add_run(
        "CRITICAL EXAMPLE:\n"
        "- Karaoke is not mentioned in this policy (implicitly permitted globally)\n"
        "- APAC addendum explicitly prohibits karaoke in APAC\n"
        "- Therefore: Karaoke is permitted globally but PROHIBITED in APAC\n"
        "- Similarly: Karaoke is permitted in Germany (not in APAC scope)"
    )
    example_run.font.bold = True

    # Documentation
    doc.add_paragraph()
    add_section_heading(doc, "6", "DOCUMENTATION AND RECEIPT REQUIREMENTS")
    doc.add_paragraph("Section 6.1 Required Documentation")
    doc.add_paragraph(
        "All expenses must include:\n"
        "- Receipt or invoice\n"
        "- Business purpose (2-3 sentences minimum)\n"
        "- Client name and organization\n"
        "- Date and location of entertainment\n"
        "- All attendees (names and titles)\n"
        "- Employee name and department",
        style='List Bullet'
    )

    # Approval
    doc.add_paragraph()
    doc.add_paragraph("Policy Approval:")
    doc.add_paragraph("Global CFO: Robert Chen")
    doc.add_paragraph("Global Chief Compliance Officer: Sarah Mitchell")
    doc.add_paragraph("CEO: Margaret Williams")

    doc.add_paragraph()
    doc.add_paragraph("Effective Date: January 1, 2025")
    doc.add_paragraph("Next Review: January 1, 2026")

    return doc

def main():
    """Create all three test documents"""
    print("Creating test documents for hallucination prevention testing...\n")

    # Create documents
    print("Creating Document 1: Global Entertainment & Client Relations Policy...")
    doc1 = create_document_1()
    doc1_path = "/tmp/Global_Entertainment_Client_Relations_Policy.docx"
    doc1.save(doc1_path)
    print(f"✓ Saved: {doc1_path}")

    print("\nCreating Document 2: Asia-Pacific Regional Addendum...")
    doc2 = create_document_2()
    doc2_path = "/tmp/Regional_Addendum_APAC_High_Risk_Activities.docx"
    doc2.save(doc2_path)
    print(f"✓ Saved: {doc2_path}")

    print("\nCreating Document 3: Global Business Travel & Expenses Policy...")
    doc3 = create_document_3()
    doc3_path = "/tmp/Global_Business_Travel_Entertainment_Expenses_Policy.docx"
    doc3.save(doc3_path)
    print(f"✓ Saved: {doc3_path}")

    print("\n" + "="*70)
    print("TEST DOCUMENTS CREATED SUCCESSFULLY")
    print("="*70)
    print(f"\nDocument 1: {doc1_path}")
    print(f"Document 2: {doc2_path}")
    print(f"Document 3: {doc3_path}")
    print("\nNext Steps:")
    print("1. Upload these documents to the compliance frontend:")
    print("   URL: https://intuition-lab.vercel.app/compliance")
    print("\n2. After upload, test with these questions:")
    print("   a) 'I'm taking a client to Karaoke in Germany. What's the risk?'")
    print("   b) 'I'm taking a client to Karaoke in Japan. What's the risk?'")
    print("   c) 'Can I take a client to a nightclub in China?'")
    print("   d) 'Entertainment expense in Vietnam - Karaoke bar. Approval needed?'")
    print("\n3. Expected Results:")
    print("   a) Germany Karaoke → LOW (green) - not in APAC scope")
    print("   b) Japan Karaoke → CRITICAL (red) - prohibited in APAC")
    print("   c) China Nightclub → CRITICAL (red) - prohibited in APAC")
    print("   d) Vietnam Karaoke → CRITICAL (red) - prohibited in APAC")
    print("\n" + "="*70)

if __name__ == "__main__":
    main()
